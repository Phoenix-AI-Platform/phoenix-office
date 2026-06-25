"""DOCX proposal renderer."""

from __future__ import annotations

from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.section import Section
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from phoenix_office.generators.proposal import ProposalFields, ProposalGenerator
from phoenix_office.models.proposal import ProposalInput

type PathLike = str | Path


class DocxProposalRenderer:
    """Render :class:`ProposalInput` into a DOCX template.

    Templates should contain placeholders using ``{{field_name}}`` tokens, for
    example ``{{customer_name}}`` or ``{{scope_block}}``.
    """

    def __init__(self, generator: ProposalGenerator | None = None) -> None:
        self.generator = generator or ProposalGenerator()

    def render(
        self,
        proposal: ProposalInput,
        template_path: PathLike,
        output_path: PathLike,
    ) -> Path:
        """Render a proposal into a DOCX file and return the saved output path."""
        fields = self.generator.prepare(proposal)
        return self.render_fields(fields, template_path=template_path, output_path=output_path)

    def render_fields(
        self,
        fields: ProposalFields,
        template_path: PathLike,
        output_path: PathLike,
    ) -> Path:
        """Render prepared proposal fields into a DOCX template."""
        document = Document(str(template_path))
        replacements = self._build_replacements(fields)
        self._replace_in_document(document, replacements)

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output))
        return output

    def _build_replacements(self, fields: ProposalFields) -> dict[str, str]:
        return {f"{{{{{name}}}}}": value for name, value in asdict(fields).items()}

    def _replace_in_document(
        self,
        document: DocumentType,
        replacements: dict[str, str],
    ) -> None:
        for paragraph in self._iter_paragraphs(document):
            self._replace_in_paragraph(paragraph, replacements)

        for section in document.sections:
            for container in self._iter_section_containers(section):
                for paragraph in self._iter_paragraphs(container):
                    self._replace_in_paragraph(paragraph, replacements)

    def _iter_section_containers(self, section: Section) -> tuple[_Cell | object, ...]:
        return (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )

    def _iter_paragraphs(self, container: DocumentType | _Cell | object):
        yield from container.paragraphs
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from self._iter_paragraphs(cell)

    def _replace_in_paragraph(self, paragraph: Paragraph, replacements: dict[str, str]) -> None:
        if not paragraph.runs:
            return

        for run in paragraph.runs:
            text = run.text
            for placeholder, value in replacements.items():
                text = text.replace(placeholder, value)
            run.text = text

        full_text = "".join(run.text for run in paragraph.runs)
        updated_text = full_text
        for placeholder, value in replacements.items():
            updated_text = updated_text.replace(placeholder, value)

        if updated_text == full_text:
            return

        paragraph.runs[0].text = updated_text
        for run in paragraph.runs[1:]:
            run.text = ""
