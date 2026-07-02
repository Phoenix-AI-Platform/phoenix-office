"""Tests for A-1 proposal intake DOCX generation CLI command."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from phoenix_office.cli import main
from phoenix_office.models.records import CustomerRecord
from phoenix_office.records import customer_record_to_json

ROOT = Path(__file__).parents[1]
A1_INTAKE_JSON = ROOT / "examples" / "proposals" / "a1_residential_tank_removal_intake.json"
A1_TEMPLATE = ROOT / "tests" / "fixtures" / "templates" / "a1_proposal_template.docx"


def read_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_cli_proposal_generate_from_intake_example(tmp_path: Path, capsys) -> None:
    output_path = tmp_path / "a1_intake_proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(A1_INTAKE_JSON),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    text = read_docx_text(output_path)

    assert exit_code == 0
    assert "Generated proposal DOCX" in captured.out
    assert str(output_path) in captured.out
    assert captured.err == ""
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert "Jane Customer" in text
    assert "123 Main St." in text
    assert "Milwaukee, WI 53202" in text
    assert "July 1, 2026" in text
    assert "Removal of 1,000 Gallon Aboveground Storage Tank" in text
    assert "TOTAL: Starting at $3,000.00" in text
    assert "Customer is responsible for access to the tank area." in text


def test_cli_proposal_generate_from_intake_invalid_intake_fails(
    tmp_path: Path, capsys
) -> None:
    input_path = tmp_path / "invalid_a1_intake.json"
    output_path = tmp_path / "a1_intake_proposal.docx"
    input_path.write_text(
        json.dumps(
            {
                "customer_name": "Jane Customer",
                "job_address": {
                    "street_address": "123 Main St.",
                    "city_state_zip": "Milwaukee, WI 53202",
                },
                "proposal_date": "2026-07-01",
                "item_description": "Removal of 1,000 Gallon Aboveground Storage Tank",
                "scope_notes": ["Remove tank"],
                "pricing_lines": [],
            }
        ),
        encoding="utf-8",
    )

    exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(input_path),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    assert exit_code == 1
    assert captured.out == ""
    assert "Error: invalid A-1 proposal intake JSON:" in captured.err
    assert "pricing_lines" in captured.err
    assert not output_path.exists()


def _write_customer_record(path: Path) -> None:
    customer = CustomerRecord(
        customer_id="customer-abby-hill",
        display_name="Abby Hill",
        phone="555-0100",
        email="abby@example.com",
        billing_street_address="999 Billing Rd.",
        billing_city_state_zip="Billing, WI 53000",
        job_street_address="W3064 Piper Rd.",
        job_city_state_zip="Whitewater, WI 53190",
    )
    path.write_text(customer_record_to_json(customer), encoding="utf-8")


def _write_customer_backed_placeholder_intake(
    tmp_path: Path, capsys
) -> tuple[Path, str]:
    customer_path = tmp_path / "abby_customer.json"
    intake_path = tmp_path / "abby_placeholder_intake.json"
    _write_customer_record(customer_path)

    exit_code = main(
        [
            "proposal",
            "customer-draft-json",
            str(customer_path),
            str(intake_path),
            "--proposal-date",
            "2026-07-01",
            "--item-description",
            "TODO: Replace with explicit item description.",
            "--scope-note",
            "Remove and dispose of tank and residual contents",
            "--pricing-description",
            "Residential tank removal",
            "--pricing-amount",
            "3000.00",
            "--special-note",
            "Customer is responsible for access to the tank area.",
        ]
    )
    output = capsys.readouterr()

    assert exit_code == 0
    assert output.err == ""
    assert intake_path.exists()
    return intake_path, "TODO: Replace with explicit item description."


def test_cli_proposal_generate_from_customer_backed_intake_blocks_placeholders(
    tmp_path: Path, capsys
) -> None:
    intake_path, _placeholder = _write_customer_backed_placeholder_intake(
        tmp_path, capsys
    )
    output_path = tmp_path / "blocked_placeholder_proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(intake_path),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    assert exit_code == 1
    assert captured.out == ""
    assert "unresolved placeholder text in A-1 proposal intake" in captured.err
    assert "Use --allow-placeholder-intake to generate anyway." in captured.err
    assert "item_description" in captured.err
    assert not output_path.exists()


def test_cli_proposal_generate_from_customer_backed_intake_allows_placeholder_override(
    tmp_path: Path, capsys
) -> None:
    intake_path, placeholder = _write_customer_backed_placeholder_intake(
        tmp_path, capsys
    )
    output_path = tmp_path / "allowed_placeholder_proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(intake_path),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
            "--allow-placeholder-intake",
        ]
    )

    captured = capsys.readouterr()
    text = read_docx_text(output_path)
    assert exit_code == 0
    assert "Generated proposal DOCX" in captured.out
    assert captured.err == ""
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert "Abby Hill" in text
    assert placeholder in text
