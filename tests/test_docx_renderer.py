"""Tests for DOCX proposal rendering."""

from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from phoenix_office.models.proposal import PricingLine, ProposalInput, ScopeItem
from phoenix_office.renderers.docx_proposal import DocxProposalRenderer

A1_TEMPLATE_PATH = Path(__file__).parent / "fixtures" / "templates" / "a1_proposal_template.docx"
PROPOSAL_PLACEHOLDERS = [
    "{{customer_name}}",
    "{{street_address}}",
    "{{city_state_zip}}",
    "{{proposal_date}}",
    "{{item_description}}",
    "{{scope_block}}",
    "{{total_line}}",
    "{{pricing_note}}",
    "{{notes}}",
]
PRICING_NOTE = (
    "Price is based on normal pumping, cleaning, and removal of the tank. "
    "Additional charges may apply depending on the quantity and condition of the tank "
    "contents or if hazardous materials are encountered."
)


def build_generic_proposal() -> ProposalInput:
    return ProposalInput(
        customer_name="Test Customer",
        street_address="123 Main St.",
        city_state_zip="Milwaukee, WI 53202",
        proposal_date=date(2026, 6, 25),
        item_description="Tank removal test item",
        scope_items=[
            ScopeItem(number=1, description="Pump contents"),
            ScopeItem(number=2, description="Remove tank"),
        ],
        pricing=PricingLine(
            amount=Decimal("1500.00"),
            is_starting_at=False,
            pricing_note="Additional charges may apply.",
        ),
        notes=["Please call before arrival.", "Gate code is 1234."],
    )


def build_abby_hill_proposal() -> ProposalInput:
    return ProposalInput(
        customer_name="Abby Hill",
        street_address="W3064 Piper Rd.",
        city_state_zip="Whitewater, WI",
        proposal_date=date(2026, 6, 25),
        item_description="Removal of 1,000 Gallon Aboveground Storage Tank",
        scope_items=[
            ScopeItem(number=1, description="Pump contents of tank (contents unknown)"),
            ScopeItem(number=2, description="Open and clean tank"),
            ScopeItem(number=3, description="Remove 1,000 gallon AST"),
            ScopeItem(number=4, description="Remove and dispose of tank and residual contents"),
        ],
        pricing=PricingLine(
            amount=Decimal("3000.00"),
            is_starting_at=True,
            pricing_note=PRICING_NOTE,
        ),
    )


def build_simple_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("Customer: {{customer_name}}")
    document.add_paragraph("Address: {{street_address}}")
    document.add_paragraph("City/State/ZIP: {{city_state_zip}}")
    document.add_paragraph("Date: {{proposal_date}}")
    document.add_paragraph("Item: {{item_description}}")
    document.add_paragraph("Scope: {{scope_block}}")
    document.add_paragraph("{{total_line}}")
    document.add_paragraph("Terms: {{pricing_note}}")
    document.add_paragraph("Notes: {{notes}}")

    split_paragraph = document.add_paragraph("Split customer: ")
    split_paragraph.add_run("{{customer_")
    split_paragraph.add_run("name}}")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table customer: {{customer_name}}"

    document.save(path)


def read_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def find_paragraph(document: Document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def render_simple_template(tmp_path: Path) -> str:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "proposal.docx"
    build_simple_template(template_path)

    DocxProposalRenderer().render(build_generic_proposal(), template_path, output_path)

    return read_docx_text(output_path)


def test_render_creates_docx_file_from_simple_template(tmp_path):
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output" / "proposal.docx"
    build_simple_template(template_path)

    result = DocxProposalRenderer().render(
        build_generic_proposal(),
        template_path,
        output_path,
    )

    assert result == output_path
    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_replaces_placeholders_in_normal_paragraphs(tmp_path):
    text = render_simple_template(tmp_path)

    assert "Customer: Test Customer" in text
    assert "Address: 123 Main St." in text
    assert "City/State/ZIP: Milwaukee, WI 53202" in text
    assert "Date: June 25, 2026" in text
    assert "Item: Tank removal test item" in text
    assert "Scope: 1. Pump contents" in text
    assert "2. Remove tank" in text
    assert "TOTAL: $1,500.00" in text
    assert "Terms: Additional charges may apply." in text


def test_replaces_placeholder_split_across_word_runs(tmp_path):
    text = render_simple_template(tmp_path)

    assert "Split customer: Test Customer" in text
    assert "{{customer_" not in text
    assert "name}}" not in text


def test_replaces_placeholders_inside_table_cell(tmp_path):
    text = render_simple_template(tmp_path)

    assert "Table customer: Test Customer" in text


def test_general_notes_render_correctly(tmp_path):
    text = render_simple_template(tmp_path)

    assert "Notes: Please call before arrival.\nGate code is 1234." in text


def test_unresolved_proposal_placeholders_are_not_left_in_generated_output(tmp_path):
    text = render_simple_template(tmp_path)

    for placeholder in PROPOSAL_PLACEHOLDERS:
        assert placeholder not in text


def test_render_creates_verified_abby_hill_docx(tmp_path):
    output_path = tmp_path / "abby_hill_proposal.docx"

    result = DocxProposalRenderer().render(
        build_abby_hill_proposal(),
        A1_TEMPLATE_PATH,
        output_path,
    )

    assert result == output_path
    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_abby_hill_text_appears_in_generated_docx(tmp_path):
    output_path = tmp_path / "abby_hill_proposal.docx"

    DocxProposalRenderer().render(build_abby_hill_proposal(), A1_TEMPLATE_PATH, output_path)

    text = read_docx_text(output_path)
    expected_text = [
        "A-1 Tank Removal LLC",
        "W178 N6006 Prairie Sky Court",
        "Menomonee Falls, Wisconsin 53051",
        "(262) 252-4030",
        "A-1 Tank Removal Proposal",
        "Date: June 25, 2026",
        "Prepared For:",
        "Abby Hill",
        "W3064 Piper Rd.",
        "Whitewater, WI",
        "Removal of 1,000 Gallon Aboveground Storage Tank",
        "1. Pump contents of tank (contents unknown)",
        "2. Open and clean tank",
        "3. Remove 1,000 gallon AST",
        "4. Remove and dispose of tank and residual contents",
        "TOTAL: Starting at $3,000.00",
        PRICING_NOTE,
        "Payment due upon completion",
        "Proposal Accepted By:",
    ]

    for expected in expected_text:
        assert expected in text
    for placeholder in PROPOSAL_PLACEHOLDERS:
        assert placeholder not in text


def test_a1_template_preserves_page_setup_content_order_and_placeholders():
    document = Document(A1_TEMPLATE_PATH)
    section = document.sections[0]

    assert section.orientation == WD_ORIENT.PORTRAIT
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1.25)
    assert section.right_margin == Inches(1.25)
    assert section.top_margin == Inches(1)
    assert section.bottom_margin == Inches(1)

    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    required_content = [
        "A-1 Tank Removal LLC",
        "A-1 Tank Removal Proposal",
        "Date: {{proposal_date}}",
        "Prepared For:",
        "{{customer_name}}",
        "{{street_address}}",
        "{{city_state_zip}}",
        "{{item_description}}",
        "Scope of Work",
        "{{scope_block}}",
        "{{total_line}}",
        "{{pricing_note}}",
        "{{notes}}",
        "Payment due upon completion",
        "Proposal Accepted By:",
        "Signature: ________________________________",
    ]

    assert not any(not text for text in paragraph_text)
    assert [paragraph_text.index(text) for text in required_content] == sorted(
        paragraph_text.index(text) for text in required_content
    )
    for placeholder in PROPOSAL_PLACEHOLDERS:
        assert sum(placeholder in text for text in paragraph_text) == 1


def test_a1_template_has_explicit_typography_spacing_and_pagination_controls():
    document = Document(A1_TEMPLATE_PATH)

    for paragraph in document.paragraphs:
        assert paragraph.paragraph_format.space_before == Pt(0)
        assert paragraph.paragraph_format.space_after is not None
        assert paragraph.paragraph_format.line_spacing == 1.0
        assert paragraph.paragraph_format.widow_control is True
        for run in paragraph.runs:
            assert run.font.name == "Times New Roman"
            assert run.font.size is not None

    company_block = [
        "A-1 Tank Removal LLC",
        "W178 N6006 Prairie Sky Court",
        "Menomonee Falls, Wisconsin 53051",
        "(262) 252-4030",
    ]
    for text in company_block:
        paragraph = find_paragraph(document, text)
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert all(run.font.size == Pt(18) for run in paragraph.runs if run.text)

    title = find_paragraph(document, "A-1 Tank Removal Proposal")
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert all(run.font.size == Pt(18) for run in title.runs if run.text)

    proposal_date = find_paragraph(document, "Date: {{proposal_date}}")
    assert proposal_date.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    prepared_for = find_paragraph(document, "Prepared For:")
    assert prepared_for.paragraph_format.left_indent == Inches(0.5)
    assert prepared_for.paragraph_format.keep_with_next is True
    for placeholder in ("{{customer_name}}", "{{street_address}}", "{{city_state_zip}}"):
        assert find_paragraph(document, placeholder).paragraph_format.left_indent == Inches(1)

    item = find_paragraph(document, "{{item_description}}")
    borders = item._p.pPr.find(qn("w:pBdr"))
    assert borders is not None
    assert borders.find(qn("w:top")) is not None
    assert borders.find(qn("w:bottom")) is not None

    assert find_paragraph(document, "Scope of Work").paragraph_format.keep_with_next is True
    scope = find_paragraph(document, "{{scope_block}}")
    assert scope.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert scope.paragraph_format.left_indent == Inches(0)
    assert scope.paragraph_format.keep_together is False

    total = find_paragraph(document, "{{total_line}}")
    assert total.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert total.paragraph_format.keep_with_next is True
    acceptance = find_paragraph(document, "Proposal Accepted By:")
    assert acceptance.paragraph_format.keep_with_next is True


def test_a1_template_and_rendered_total_line_are_completely_bold(tmp_path):
    template = Document(A1_TEMPLATE_PATH)
    template_total = find_paragraph(template, "{{total_line}}")
    assert template_total.runs
    assert all(run.bold is True for run in template_total.runs if run.text)

    output_path = tmp_path / "proposal.docx"
    DocxProposalRenderer().render(build_generic_proposal(), A1_TEMPLATE_PATH, output_path)

    rendered = Document(output_path)
    rendered_total = next(
        paragraph for paragraph in rendered.paragraphs if paragraph.text.startswith("TOTAL:")
    )
    assert rendered_total.text == "TOTAL: $1,500.00"
    assert rendered_total.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert all(run.bold is True for run in rendered_total.runs if run.text)
    for placeholder in PROPOSAL_PLACEHOLDERS:
        assert placeholder not in read_docx_text(output_path)
