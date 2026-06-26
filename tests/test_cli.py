"""Tests for the Phoenix Office command-line interface."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from phoenix_office.cli import main

ROOT = Path(__file__).parents[1]
EXAMPLE_JSON = ROOT / "examples" / "proposals" / "abby_hill.json"
A1_TEMPLATE = ROOT / "tests" / "fixtures" / "templates" / "a1_proposal_template.docx"


def read_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_cli_generates_abby_hill_docx_from_example_json(tmp_path, capsys):
    output_path = tmp_path / "abby_hill_proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate",
            str(EXAMPLE_JSON),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    text = read_docx_text(output_path)
    assert exit_code == 0
    assert "Generated proposal DOCX" in captured.out
    assert str(output_path) in captured.out
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert "Abby Hill" in text
    assert "W3064 Piper Rd." in text
    assert "Whitewater, WI" in text
    assert "June 25, 2026" in text
    assert "Removal of 1,000 Gallon Aboveground Storage Tank" in text
    assert "TOTAL: Starting at $3,000.00" in text


def test_cli_intake_generates_docx_and_optional_json(tmp_path, capsys, monkeypatch):
    output_path = tmp_path / "intake_proposal.docx"
    json_output_path = tmp_path / "intake_proposal.json"
    answers = iter(
        [
            "Jane Customer",
            "123 Main St.",
            "Milwaukee, WI 53202",
            "2026-06-25",
            "1,000",
            "AST",
            "unknown",
            "starting-at",
            "3000",
            "Additional charges may apply after inspection.",
            "Please call before arrival.",
        ]
    )
    monkeypatch.setattr("builtins.input", lambda _message: next(answers))

    exit_code = main(
        [
            "proposal",
            "intake",
            "--template",
            str(A1_TEMPLATE),
            "--output",
            str(output_path),
            "--json-output",
            str(json_output_path),
        ]
    )

    captured = capsys.readouterr()
    text = read_docx_text(output_path)
    proposal_json = json.loads(json_output_path.read_text(encoding="utf-8"))

    assert exit_code == 0
    assert "Generated proposal DOCX" in captured.out
    assert "Wrote proposal JSON" in captured.out
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert json_output_path.exists()
    assert proposal_json["customer_name"] == "Jane Customer"
    assert proposal_json["pricing"]["is_starting_at"] is True
    assert proposal_json["pricing"]["pricing_note"] == (
        "Additional charges may apply after inspection."
    )
    assert "Jane Customer" in text
    assert "Remove one 1,000 gallon AST tank located at 123 Main St., " in text
    assert "Milwaukee, WI 53202 (contents unknown)." in text
    assert "TOTAL: Starting at $3,000.00" in text
    assert "Additional charges may apply after inspection." in text
    assert "Please call before arrival." in text


def test_cli_intake_preset_flow_generates_expected_docx_text(tmp_path, capsys, monkeypatch):
    output_path = tmp_path / "preset_intake_proposal.docx"
    answers = iter(
        [
            "Jane Customer",
            "123 Main St.",
            "Milwaukee, WI 53202",
            "2026-06-25",
            "1",
            "AST",
            "unknown",
            "starting-at",
            "3000",
            "3",
            "1",
        ]
    )
    monkeypatch.setattr("builtins.input", lambda _message: next(answers))

    exit_code = main(
        [
            "proposal",
            "intake",
            "--template",
            str(A1_TEMPLATE),
            "--output",
            str(output_path),
        ]
    )

    captured = capsys.readouterr()
    text = read_docx_text(output_path)

    assert exit_code == 0
    assert "Generated proposal DOCX" in captured.out
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert "Jane Customer" in text
    assert "123 Main St." in text
    assert "Milwaukee, WI 53202" in text
    assert (
        "Remove one 275 gallon AST tank located at 123 Main St., "
        "Milwaukee, WI 53202 (contents unknown)."
    ) in text
    assert "TOTAL: Starting at $3,000.00" in text
    assert "Starting at price based on information provided prior to inspection." in text
    assert "Payment due upon completion." in text


def test_cli_creates_output_directory(tmp_path):
    output_path = tmp_path / "nested" / "output" / "abby_hill_proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate",
            str(EXAMPLE_JSON),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    assert exit_code == 0
    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_cli_fails_cleanly_when_input_json_is_missing(tmp_path, capsys):
    missing_json = tmp_path / "missing.json"
    output_path = tmp_path / "proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate",
            str(missing_json),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    assert exit_code != 0
    assert "JSON input file does not exist" in captured.err
    assert not output_path.exists()


def test_cli_fails_cleanly_when_template_path_is_missing(tmp_path, capsys):
    missing_template = tmp_path / "missing-template.docx"
    output_path = tmp_path / "proposal.docx"

    exit_code = main(
        [
            "proposal",
            "generate",
            str(EXAMPLE_JSON),
            str(output_path),
            "--template",
            str(missing_template),
        ]
    )

    captured = capsys.readouterr()
    assert exit_code != 0
    assert "DOCX template file does not exist" in captured.err
    assert not output_path.exists()


def test_cli_fails_cleanly_when_json_is_invalid(tmp_path, capsys):
    invalid_json = tmp_path / "invalid.json"
    output_path = tmp_path / "proposal.docx"
    invalid_json.write_text("{not valid json", encoding="utf-8")

    exit_code = main(
        [
            "proposal",
            "generate",
            str(invalid_json),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    assert exit_code != 0
    assert "Invalid JSON" in captured.err
    assert not output_path.exists()


def test_cli_fails_cleanly_when_model_validation_fails(tmp_path, capsys):
    invalid_proposal = tmp_path / "invalid-proposal.json"
    output_path = tmp_path / "proposal.docx"
    invalid_proposal.write_text(
        "{\n"
        '  "customer_name": "",\n'
        '  "street_address": "123 Main St.",\n'
        '  "city_state_zip": "Milwaukee, WI",\n'
        '  "proposal_date": "2026-06-25",\n'
        '  "item_description": "Tank removal",\n'
        '  "scope_items": [],\n'
        '  "pricing": {"amount": "0"}\n'
        "}\n",
        encoding="utf-8",
    )

    exit_code = main(
        [
            "proposal",
            "generate",
            str(invalid_proposal),
            str(output_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )

    captured = capsys.readouterr()
    assert exit_code != 0
    assert "Invalid proposal input" in captured.err
    assert not output_path.exists()
