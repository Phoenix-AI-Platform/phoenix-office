"""End-to-end smoke coverage for the A-1 proposal intake CLI workflow."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from phoenix_office.cli import main
from phoenix_office.models.records import CustomerRecord
from phoenix_office.records import customer_record_to_json

ROOT = Path(__file__).parents[1]
A1_TEMPLATE = ROOT / "tests" / "fixtures" / "templates" / "a1_proposal_template.docx"


def read_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_cli_a1_intake_workflow_smoke(tmp_path: Path, capsys) -> None:
    intake_path = tmp_path / "a1_proposal_intake.json"
    proposal_input_path = tmp_path / "a1_proposal_input.json"
    proposal_docx_path = tmp_path / "a1_proposal.docx"

    draft_exit_code = main(["proposal", "draft-json", str(intake_path)])
    draft_output = capsys.readouterr()

    assert draft_exit_code == 0
    assert "Wrote proposal draft JSON" in draft_output.out
    assert draft_output.err == ""
    assert intake_path.exists()

    validate_exit_code = main(["proposal", "intake-validate", str(intake_path)])
    validate_output = capsys.readouterr()

    assert validate_exit_code == 0
    assert "A-1 proposal intake validation passed" in validate_output.out
    assert validate_output.err == ""

    inspect_exit_code = main(["proposal", "intake-inspect", str(intake_path), "--json"])
    inspect_output = capsys.readouterr()
    inspect_payload = json.loads(inspect_output.out)

    assert inspect_exit_code == 0
    assert inspect_payload["customer_name"] == "Jane Customer"
    assert inspect_payload["job_address"] == {
        "city_state_zip": "Milwaukee, WI 53202",
        "street_address": "123 Main St.",
    }
    assert inspect_payload["item_description"] == (
        "Removal of 1,000 Gallon Aboveground Storage Tank"
    )
    assert inspect_payload["pricing_amount"] == "3000.00"
    assert inspect_payload["scope_count"] == 4
    assert inspect_output.err == ""

    normalize_exit_code = main(
        ["proposal", "intake-normalize", str(intake_path), str(proposal_input_path)]
    )
    normalize_output = capsys.readouterr()

    assert normalize_exit_code == 0
    assert "Normalized proposal intake JSON" in normalize_output.out
    assert normalize_output.err == ""
    assert proposal_input_path.exists()

    generate_exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(intake_path),
            str(proposal_docx_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )
    generate_output = capsys.readouterr()

    assert generate_exit_code == 0
    assert "Generated proposal DOCX" in generate_output.out
    assert generate_output.err == ""
    assert proposal_docx_path.exists()
    assert proposal_docx_path.stat().st_size > 0

    docx_text = read_docx_text(proposal_docx_path)
    assert "Jane Customer" in docx_text
    assert "123 Main St." in docx_text
    assert "Milwaukee, WI 53202" in docx_text
    assert "July 1, 2026" in docx_text
    assert "Removal of 1,000 Gallon Aboveground Storage Tank" in docx_text
    assert "TOTAL: Starting at $3,000.00" in docx_text
    assert "Customer is responsible for access to the tank area." in docx_text


def _write_customer_record(path: Path) -> None:
    customer = CustomerRecord(
        customer_id="customer-abby-hill",
        display_name="Abby Hill",
        phone="555-0100",
        email="abby@example.com",
        billing_street_address="999 Billing Rd.",
        billing_city_state_zip="Billing, WI 53000",
        job_street_address="W3064 Piper Rd.",
        job_city_state_zip="Whitewater, WI 53190",
    )
    path.write_text(customer_record_to_json(customer), encoding="utf-8")


def test_cli_customer_backed_a1_intake_workflow_smoke(
    tmp_path: Path, capsys
) -> None:
    customer_path = tmp_path / "abby_customer.json"
    intake_path = tmp_path / "abby_intake.json"
    proposal_input_path = tmp_path / "abby_proposal_input.json"
    proposal_docx_path = tmp_path / "abby_proposal.docx"
    _write_customer_record(customer_path)

    draft_exit_code = main(
        [
            "proposal",
            "customer-draft-json",
            str(customer_path),
            str(intake_path),
            "--proposal-date",
            "2026-07-01",
            "--item-description",
            "Removal of 1,000 Gallon Aboveground Storage Tank",
            "--scope-note",
            "Pump contents of tank (contents unknown)",
            "--scope-note",
            "Remove and dispose of tank and residual contents",
            "--pricing-description",
            "Residential tank removal",
            "--pricing-amount",
            "3000.00",
            "--pricing-note",
            "Explicit pricing note.",
            "--starting-at",
            "--special-note",
            "Customer is responsible for access to the tank area.",
        ]
    )
    draft_output = capsys.readouterr()
    draft_payload = json.loads(intake_path.read_text(encoding="utf-8"))

    assert draft_exit_code == 0
    assert "Wrote customer-backed proposal draft JSON" in draft_output.out
    assert draft_output.err == ""
    assert draft_payload["customer_name"] == "Abby Hill"
    assert draft_payload["job_address"] == {
        "city_state_zip": "Whitewater, WI 53190",
        "street_address": "W3064 Piper Rd.",
    }
    assert draft_payload["scope_notes"] == [
        "Pump contents of tank (contents unknown)",
        "Remove and dispose of tank and residual contents",
    ]

    validate_exit_code = main(["proposal", "intake-validate", str(intake_path)])
    validate_output = capsys.readouterr()

    assert validate_exit_code == 0
    assert "A-1 proposal intake validation passed" in validate_output.out
    assert validate_output.err == ""

    inspect_exit_code = main(["proposal", "intake-inspect", str(intake_path), "--json"])
    inspect_output = capsys.readouterr()
    inspect_payload = json.loads(inspect_output.out)

    assert inspect_exit_code == 0
    assert inspect_output.err == ""
    assert inspect_payload["customer_name"] == "Abby Hill"
    assert inspect_payload["job_address"] == {
        "city_state_zip": "Whitewater, WI 53190",
        "street_address": "W3064 Piper Rd.",
    }
    assert inspect_payload["proposal_date"] == "2026-07-01"
    assert inspect_payload["item_description"] == (
        "Removal of 1,000 Gallon Aboveground Storage Tank"
    )
    assert inspect_payload["pricing_amount"] == "3000.00"
    assert inspect_payload["scope_count"] == 2

    normalize_exit_code = main(
        ["proposal", "intake-normalize", str(intake_path), str(proposal_input_path)]
    )
    normalize_output = capsys.readouterr()
    proposal_input_payload = json.loads(proposal_input_path.read_text(encoding="utf-8"))

    assert normalize_exit_code == 0
    assert "Normalized proposal intake JSON" in normalize_output.out
    assert normalize_output.err == ""
    assert proposal_input_payload["customer_name"] == "Abby Hill"
    assert proposal_input_payload["street_address"] == "W3064 Piper Rd."
    assert proposal_input_payload["city_state_zip"] == "Whitewater, WI 53190"
    assert proposal_input_payload["pricing"]["amount"] == "3000.00"
    assert proposal_input_payload["pricing"]["is_starting_at"] is True
    assert proposal_input_payload["notes"] == [
        "Customer is responsible for access to the tank area."
    ]

    generate_exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(intake_path),
            str(proposal_docx_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )
    generate_output = capsys.readouterr()

    assert generate_exit_code == 0
    assert "Generated proposal DOCX" in generate_output.out
    assert generate_output.err == ""
    assert proposal_docx_path.exists()
    assert proposal_docx_path.stat().st_size > 0

    docx_text = read_docx_text(proposal_docx_path)
    assert "Abby Hill" in docx_text
    assert "W3064 Piper Rd." in docx_text
    assert "Whitewater, WI 53190" in docx_text
    assert "July 1, 2026" in docx_text
    assert "Removal of 1,000 Gallon Aboveground Storage Tank" in docx_text
    assert "Pump contents of tank (contents unknown)" in docx_text
    assert "TOTAL: Starting at $3,000.00" in docx_text
    assert "Customer is responsible for access to the tank area." in docx_text


def test_cli_readiness_workflow_smoke(tmp_path: Path, capsys) -> None:
    intake_path = tmp_path / "a1_proposal_intake.json"
    blocked_intake_path = tmp_path / "blocked_a1_proposal_intake.json"
    proposal_input_path = tmp_path / "a1_proposal_input.json"
    blocked_proposal_input_path = tmp_path / "blocked_a1_proposal_input.json"
    blocked_docx_path = tmp_path / "blocked_a1_proposal.docx"

    draft_exit_code = main(["proposal", "draft-json", str(intake_path)])
    draft_output = capsys.readouterr()

    assert draft_exit_code == 0
    assert draft_output.err == ""
    assert intake_path.exists()

    intake_ready_exit_code = main(
        ["proposal", "intake-readiness", str(intake_path), "--json"]
    )
    intake_ready_output = capsys.readouterr()
    intake_ready_payload = json.loads(intake_ready_output.out)

    assert intake_ready_exit_code == 0
    assert intake_ready_payload == {
        "blocker_field_paths": [],
        "input_path": str(intake_path),
        "ready_for_normalization": True,
        "status": "ready",
    }
    assert intake_ready_output.err == ""

    blocked_intake_payload = json.loads(intake_path.read_text(encoding="utf-8"))
    blocked_intake_payload["item_description"] = (
        "TODO: Replace with explicit item description."
    )
    blocked_intake_payload["scope_notes"][0] = "Replace with explicit first scope note."
    blocked_intake_path.write_text(
        json.dumps(blocked_intake_payload),
        encoding="utf-8",
    )

    intake_blocked_exit_code = main(
        ["proposal", "intake-readiness", str(blocked_intake_path), "--json"]
    )
    intake_blocked_output = capsys.readouterr()
    intake_blocked_payload = json.loads(intake_blocked_output.out)

    assert intake_blocked_exit_code == 1
    assert intake_blocked_payload == {
        "blocker_field_paths": [
            "item_description",
            "scope_notes[0]",
        ],
        "input_path": str(blocked_intake_path),
        "ready_for_normalization": False,
        "status": "blocked",
    }
    assert intake_blocked_output.err == ""

    normalize_exit_code = main(
        ["proposal", "intake-normalize", str(intake_path), str(proposal_input_path)]
    )
    normalize_output = capsys.readouterr()

    assert normalize_exit_code == 0
    assert "Normalized proposal intake JSON" in normalize_output.out
    assert normalize_output.err == ""
    assert proposal_input_path.exists()

    proposal_ready_exit_code = main(
        ["proposal", "readiness", str(proposal_input_path), "--json"]
    )
    proposal_ready_output = capsys.readouterr()
    proposal_ready_payload = json.loads(proposal_ready_output.out)

    assert proposal_ready_exit_code == 0
    assert proposal_ready_payload == {
        "blocker_field_paths": [],
        "input_path": str(proposal_input_path),
        "ready_for_docx_generation": True,
        "status": "ready",
    }
    assert proposal_ready_output.err == ""

    blocked_proposal_payload = json.loads(
        proposal_input_path.read_text(encoding="utf-8")
    )
    blocked_proposal_payload["item_description"] = (
        "TODO: Replace with explicit item description."
    )
    blocked_proposal_input_path.write_text(
        json.dumps(blocked_proposal_payload),
        encoding="utf-8",
    )

    proposal_blocked_exit_code = main(
        ["proposal", "readiness", str(blocked_proposal_input_path), "--json"]
    )
    proposal_blocked_output = capsys.readouterr()
    proposal_blocked_payload = json.loads(proposal_blocked_output.out)

    assert proposal_blocked_exit_code == 1
    assert proposal_blocked_payload == {
        "blocker_field_paths": ["item_description"],
        "input_path": str(blocked_proposal_input_path),
        "ready_for_docx_generation": False,
        "status": "blocked",
    }
    assert proposal_blocked_output.err == ""

    blocked_generate_exit_code = main(
        [
            "proposal",
            "generate-from-intake",
            str(blocked_intake_path),
            str(blocked_docx_path),
            "--template",
            str(A1_TEMPLATE),
        ]
    )
    blocked_generate_output = capsys.readouterr()

    assert blocked_generate_exit_code == 1
    assert blocked_generate_output.out == ""
    assert "unresolved placeholder text in A-1 proposal intake" in (
        blocked_generate_output.err
    )
    assert "Placeholder fields: item_description, scope_notes[0]" in (
        blocked_generate_output.err
    )
    assert not blocked_docx_path.exists()
