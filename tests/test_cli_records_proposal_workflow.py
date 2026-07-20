"""Smoke test for the record-backed proposal CLI workflow."""

import hashlib
import json
import sqlite3
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document

import phoenix_office.cli as cli
import phoenix_office.proposal_build as proposal_build
import phoenix_office.records.sqlite as sqlite_records
from phoenix_office.cli import build_parser, main
from phoenix_office.models.proposal import ProposalInput
from phoenix_office.records import (
    SQLiteCustomerRepository,
    SQLiteJobRepository,
    create_sqlite_record_store,
    record_proposal_details_from_file,
)

ROOT = Path(__file__).parents[1]
CUSTOMER_EXAMPLE = ROOT / "examples" / "records" / "customer_abby_hill.json"
JOB_EXAMPLE = ROOT / "examples" / "records" / "job_abby_hill.json"
DETAILS_EXAMPLE = ROOT / "examples" / "records" / "proposal_details_abby_hill.json"
TEMPLATE = ROOT / "tests" / "fixtures" / "templates" / "a1_proposal_template.docx"
SAMPLE_JOB_EXAMPLE = ROOT / "examples" / "records" / "job_sample_north_prairie.json"


def _collapse_whitespace(text: str) -> str:
    return " ".join(text.split())


def _import_abby_records(db_path: Path) -> None:
    assert main(
        ["records", "import", "customer", str(CUSTOMER_EXAMPLE), "--db", str(db_path)]
    ) == 0
    assert main(
        ["records", "import", "job", str(JOB_EXAMPLE), "--db", str(db_path)]
    ) == 0


def _proposal_build_args(
    db_path: Path,
    output_json: Path,
    output_docx: Path,
    *,
    customer_id: str = "customer-abby-hill",
    job_id: str = "job-abby-hill",
    details_path: Path = DETAILS_EXAMPLE,
    template_path: Path = TEMPLATE,
) -> list[str]:
    return [
        "records",
        "proposal-build",
        customer_id,
        job_id,
        str(details_path),
        str(output_json),
        str(output_docx),
        "--db",
        str(db_path),
        "--template",
        str(template_path),
    ]


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    return _collapse_whitespace(" ".join(values))


def _assert_no_outputs(output_json: Path, output_docx: Path) -> None:
    assert not output_json.exists()
    assert not output_docx.exists()


def _database_snapshot(db_path: Path) -> tuple[object, ...]:
    database_uri = f"{db_path.resolve().as_uri()}?mode=ro&immutable=1"
    with sqlite3.connect(database_uri, uri=True) as connection:
        sqlite_master = tuple(
            connection.execute(
                """
                SELECT type, name, tbl_name, sql
                FROM sqlite_master
                ORDER BY type, name, tbl_name
                """
            ).fetchall()
        )
        table_names = {
            row[1] for row in sqlite_master if row[0] == "table" and row[1] != "sqlite_sequence"
        }
        table_contents = tuple(
            (table_name, tuple(connection.execute(f"SELECT * FROM {table_name} ORDER BY rowid")))
            for table_name in sorted(table_names)
        )
    database_bytes = db_path.read_bytes()
    return (
        hashlib.sha256(database_bytes).hexdigest(),
        len(database_bytes),
        sqlite_master,
        table_contents,
    )


def _assert_database_has_no_sidecars_or_outputs(tmp_path: Path, db_path: Path) -> None:
    assert set(tmp_path.iterdir()) == {db_path}


def _assert_no_sqlite_sidecars(db_path: Path) -> None:
    for suffix in ("-wal", "-shm", "-journal"):
        assert not Path(f"{db_path}{suffix}").exists()


def _sqlite_sidecar_snapshot(db_path: Path) -> dict[str, bytes | None]:
    return {
        suffix: path.read_bytes() if path.exists() else None
        for suffix in ("-wal", "-shm", "-journal")
        for path in (Path(f"{db_path}{suffix}"),)
    }


def _enable_persistent_wal_mode(db_path: Path) -> None:
    with sqlite3.connect(db_path) as connection:
        result = connection.execute("PRAGMA journal_mode=WAL").fetchone()
        assert result is not None
        assert result[0].lower() == "wal"
    _assert_no_sqlite_sidecars(db_path)


def _write_details(tmp_path: Path, transform) -> Path:
    payload = json.loads(DETAILS_EXAMPLE.read_text(encoding="utf-8"))
    transform(payload)
    path = tmp_path / "proposal_details.json"
    path.write_text(json.dumps(payload), encoding="utf-8")
    return path


def test_records_cli_proposal_input_to_docx_workflow(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    proposal_input_path = tmp_path / "proposal" / "abby_hill_proposal_input.json"
    output_docx_path = tmp_path / "proposal" / "abby_hill_proposal.docx"

    assert main(
        [
            "records",
            "import",
            "customer",
            str(CUSTOMER_EXAMPLE),
            "--db",
            str(db_path),
        ]
    ) == 0
    capsys.readouterr()

    assert main(
        ["records", "import", "job", str(JOB_EXAMPLE), "--db", str(db_path)]
    ) == 0
    capsys.readouterr()

    assert main(
        [
            "records",
            "proposal-details",
            "validate",
            str(DETAILS_EXAMPLE),
        ]
    ) == 0
    capsys.readouterr()

    assert main(
        [
            "records",
            "proposal-input",
            "customer-abby-hill",
            "job-abby-hill",
            str(DETAILS_EXAMPLE),
            str(proposal_input_path),
            "--db",
            str(db_path),
        ]
    ) == 0
    capsys.readouterr()

    assert proposal_input_path.exists()
    proposal = ProposalInput.model_validate(
        json.loads(proposal_input_path.read_text(encoding="utf-8"))
    )
    assert proposal.customer_name == "Abby Hill"
    assert (
        proposal.item_description
        == "Removal of 1,000 Gallon Aboveground Storage Tank"
    )
    assert proposal.street_address == "123 Main St."
    assert proposal.city_state_zip == "Menomonee Falls, WI 53051"
    assert len(proposal.scope_items) == 4
    assert [item.description for item in proposal.scope_items] == [
        "Pump contents of tank (contents unknown)",
        "Open and clean tank",
        "Remove 1,000 gallon AST",
        "Remove and dispose of tank and residual contents",
    ]
    assert proposal.pricing.amount == Decimal("3000.00")
    assert proposal.pricing.is_starting_at is True
    assert proposal.notes == []

    assert main(["proposal", "validate", str(proposal_input_path)]) == 0
    capsys.readouterr()

    assert main(["proposal", "inspect", str(proposal_input_path)]) == 0
    captured = capsys.readouterr()
    inspect_output = _collapse_whitespace(captured.out)
    assert "Customer: Abby Hill" in inspect_output
    assert "Site Address: 123 Main St., Menomonee Falls, WI 53051" in inspect_output
    assert "Item Description: Removal of 1,000 Gallon Aboveground Storage Tank" in inspect_output
    assert "Scope Items: 4" in inspect_output
    assert "Total: Starting at $3,000.00" in inspect_output
    assert "Notes: none" in inspect_output

    assert main(["proposal", "inspect", str(proposal_input_path), "--json"]) == 0
    captured = capsys.readouterr()
    inspect_payload = json.loads(captured.out)
    assert inspect_payload["customer_name"] == "Abby Hill"
    assert inspect_payload["street_address"] == "123 Main St."
    assert inspect_payload["city_state_zip"] == "Menomonee Falls, WI 53051"
    assert (
        inspect_payload["item_description"]
        == "Removal of 1,000 Gallon Aboveground Storage Tank"
    )
    assert [item["description"] for item in inspect_payload["scope_items"]] == [
        "Pump contents of tank (contents unknown)",
        "Open and clean tank",
        "Remove 1,000 gallon AST",
        "Remove and dispose of tank and residual contents",
    ]
    assert inspect_payload["pricing"]["amount"] == "3000.00"
    assert inspect_payload["pricing"]["is_starting_at"] is True
    assert inspect_payload["notes"] == []
    assert captured.err == ""

    assert main(
        [
            "proposal",
            "generate",
            str(proposal_input_path),
            str(output_docx_path),
            "--template",
            str(TEMPLATE),
        ]
    ) == 0
    capsys.readouterr()

    assert output_docx_path.exists()
    assert output_docx_path.stat().st_size > 0


def test_proposal_build_parser_exposes_exact_command_contract(tmp_path: Path) -> None:
    parser = build_parser()
    args = parser.parse_args(
        _proposal_build_args(
            tmp_path / "records.sqlite",
            tmp_path / "proposal.json",
            tmp_path / "proposal.docx",
        )
    )

    assert args.func is cli.build_record_proposal
    assert args.customer_id == "customer-abby-hill"
    assert args.job_id == "job-abby-hill"
    assert args.details_json == DETAILS_EXAMPLE
    assert args.output_proposal_input_json == tmp_path / "proposal.json"
    assert args.output_docx == tmp_path / "proposal.docx"
    assert args.db == tmp_path / "records.sqlite"
    assert args.template == TEMPLATE


def test_proposal_build_cli_delegates_to_in_process_service(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    db_path = tmp_path / "records.sqlite"
    db_path.write_bytes(b"delegation-only database placeholder")
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"
    details = record_proposal_details_from_file(DETAILS_EXAMPLE)
    proposal = ProposalInput(
        customer_name="Delegated Customer",
        street_address="1 Example Way",
        city_state_zip="Example, WI 00000",
        proposal_date=details.proposal_date,
        item_description=details.item_description,
        scope_items=details.scope_items,
        pricing=details.pricing,
        notes=details.notes,
        company_config=details.company_config,
    )
    expected_result = proposal_build.ProposalDraftBuildResult(
        proposal_input=proposal,
        proposal_input_json_path=output_json,
        proposal_docx_path=output_docx,
        summary_lines=("Delegated summary",),
    )
    received_requests: list[proposal_build.ProposalDraftBuildRequest] = []

    def delegate(request: proposal_build.ProposalDraftBuildRequest):
        received_requests.append(request)
        return expected_result

    monkeypatch.setattr(cli, "build_proposal_draft", delegate)

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 0
    captured = capsys.readouterr()

    assert received_requests == [
        proposal_build.ProposalDraftBuildRequest(
            customer_id="customer-abby-hill",
            job_id="job-abby-hill",
            details=details,
            database_path=db_path,
            template_path=TEMPLATE,
            proposal_input_json_output_path=output_json,
            proposal_docx_output_path=output_docx,
        )
    ]
    assert captured.out.splitlines() == [
        "Delegated summary",
        f"ProposalInput JSON: {output_json}",
        f"Proposal DOCX: {output_docx}",
    ]
    assert captured.err == ""


@pytest.mark.parametrize("positional_count", range(5))
def test_proposal_build_parser_requires_all_positionals(
    tmp_path: Path,
    positional_count: int,
) -> None:
    positionals = [
        "customer-abby-hill",
        "job-abby-hill",
        str(DETAILS_EXAMPLE),
        str(tmp_path / "proposal.json"),
        str(tmp_path / "proposal.docx"),
    ]
    argv = [
        "records",
        "proposal-build",
        *positionals[:positional_count],
        "--db",
        str(tmp_path / "records.sqlite"),
        "--template",
        str(TEMPLATE),
    ]

    with pytest.raises(SystemExit):
        build_parser().parse_args(argv)


@pytest.mark.parametrize("missing_option", ["--db", "--template"])
def test_proposal_build_parser_requires_db_and_template(
    tmp_path: Path,
    missing_option: str,
) -> None:
    argv = _proposal_build_args(
        tmp_path / "records.sqlite",
        tmp_path / "proposal.json",
        tmp_path / "proposal.docx",
    )
    option_index = argv.index(missing_option)
    del argv[option_index : option_index + 2]

    with pytest.raises(SystemExit):
        build_parser().parse_args(argv)


@pytest.mark.parametrize(
    "forbidden_option",
    ["--force", "--allow-placeholder-proposal-input", "--allow-placeholder-intake"],
)
def test_proposal_build_parser_exposes_no_bypass_options(
    tmp_path: Path,
    forbidden_option: str,
) -> None:
    argv = _proposal_build_args(
        tmp_path / "records.sqlite",
        tmp_path / "proposal.json",
        tmp_path / "proposal.docx",
    )

    with pytest.raises(SystemExit):
        build_parser().parse_args([*argv, forbidden_option])


def test_proposal_build_creates_reviewable_artifacts_without_mutating_records(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    db_path = tmp_path / "records.sqlite"
    output_dir = tmp_path / "proposal"
    output_json = output_dir / "proposal_input.json"
    output_docx = output_dir / "proposal.docx"
    _import_abby_records(db_path)
    capsys.readouterr()

    store = create_sqlite_record_store(db_path)
    customer_before = store.customers.get_customer("customer-abby-hill")
    job_before = store.jobs.get_job("job-abby-hill")

    def unexpected_call(*_args, **_kwargs):
        raise AssertionError("out-of-scope operation invoked")

    for name in (
        "import_customer_record_file",
        "import_customer_records_file",
        "import_job_record_file",
        "import_job_records_file",
        "execute_core_command",
        "execute_sdk_command",
        "send_email",
        "upload_file",
        "file_proposal",
    ):
        monkeypatch.setattr(cli, name, unexpected_call, raising=False)
    monkeypatch.setattr(cli.subprocess, "run", unexpected_call)
    monkeypatch.setattr(cli.subprocess, "Popen", unexpected_call)

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 0
    captured = capsys.readouterr()

    assert output_json.exists()
    assert output_docx.exists()
    assert output_docx.stat().st_size > 0
    proposal = ProposalInput.model_validate_json(output_json.read_text(encoding="utf-8"))
    assert proposal.customer_name == "Abby Hill"
    assert proposal.street_address == "123 Main St."
    assert proposal.city_state_zip == "Menomonee Falls, WI 53051"
    assert proposal.item_description == "Removal of 1,000 Gallon Aboveground Storage Tank"
    assert [item.description for item in proposal.scope_items] == [
        "Pump contents of tank (contents unknown)",
        "Open and clean tank",
        "Remove 1,000 gallon AST",
        "Remove and dispose of tank and residual contents",
    ]
    assert proposal.pricing.amount == Decimal("3000.00")
    assert proposal.pricing.is_starting_at is True
    assert proposal.notes == []

    docx_text = _docx_text(output_docx)
    assert "Abby Hill" in docx_text
    assert "123 Main St." in docx_text
    assert "Removal of 1,000 Gallon Aboveground Storage Tank" in docx_text
    assert "Starting at $3,000.00" in docx_text

    summary = _collapse_whitespace(captured.out)
    assert "Customer: Abby Hill" in summary
    assert "Site Address: 123 Main St., Menomonee Falls, WI 53051" in summary
    assert "Item Description: Removal of 1,000 Gallon Aboveground Storage Tank" in summary
    assert "Scope Items: 4" in summary
    assert "Pricing Lines: 1" in summary
    assert "Total: Starting at $3,000.00" in summary
    assert "Notes: none" in summary
    assert "Company: A-1 Tank Removal LLC" in summary
    assert f"ProposalInput JSON: {output_json}" in captured.out
    assert f"Proposal DOCX: {output_docx}" in captured.out
    assert captured.err == ""

    store_after = create_sqlite_record_store(db_path)
    assert store_after.customers.get_customer("customer-abby-hill") == customer_before
    assert store_after.jobs.get_job("job-abby-hill") == job_before
    assert set(output_dir.iterdir()) == {output_json, output_docx}


def test_proposal_build_repeated_fresh_outputs_are_equivalent(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    first_json = tmp_path / "first" / "proposal.json"
    first_docx = tmp_path / "first" / "proposal.docx"
    second_json = tmp_path / "second" / "proposal.json"
    second_docx = tmp_path / "second" / "proposal.docx"

    assert main(_proposal_build_args(db_path, first_json, first_docx)) == 0
    capsys.readouterr()
    assert main(_proposal_build_args(db_path, second_json, second_docx)) == 0
    capsys.readouterr()

    assert first_json.read_text(encoding="utf-8") == second_json.read_text(encoding="utf-8")
    assert _docx_text(first_docx) == _docx_text(second_docx)


def test_proposal_build_service_and_cli_outputs_are_equivalent(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    service_json = tmp_path / "service" / "proposal.json"
    service_docx = tmp_path / "service" / "proposal.docx"
    cli_json = tmp_path / "cli" / "proposal.json"
    cli_docx = tmp_path / "cli" / "proposal.docx"

    result = proposal_build.build_proposal_draft(
        proposal_build.ProposalDraftBuildRequest(
            customer_id="customer-abby-hill",
            job_id="job-abby-hill",
            details=record_proposal_details_from_file(DETAILS_EXAMPLE),
            database_path=db_path,
            template_path=TEMPLATE,
            proposal_input_json_output_path=service_json,
            proposal_docx_output_path=service_docx,
        )
    )
    assert main(_proposal_build_args(db_path, cli_json, cli_docx)) == 0
    capsys.readouterr()

    assert result.proposal_input == ProposalInput.model_validate_json(
        cli_json.read_text(encoding="utf-8")
    )
    assert service_json.read_text(encoding="utf-8") == cli_json.read_text(
        encoding="utf-8"
    )
    assert _docx_text(service_docx) == _docx_text(cli_docx)


def test_proposal_build_missing_database_does_not_create_it(tmp_path: Path, capsys) -> None:
    db_path = tmp_path / "missing.sqlite"
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    captured = capsys.readouterr()

    assert "records database does not exist" in captured.err
    assert not db_path.exists()
    _assert_no_outputs(output_json, output_docx)


def test_proposal_build_rejects_database_directory(tmp_path: Path, capsys) -> None:
    db_path = tmp_path / "records"
    db_path.mkdir()
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    assert "records database path is not a file" in capsys.readouterr().err
    _assert_no_outputs(output_json, output_docx)


@pytest.mark.parametrize(
    ("customer_id", "job_id", "expected_error"),
    [
        ("missing-customer", "job-abby-hill", "Customer not found"),
        ("customer-abby-hill", "missing-job", "Job not found"),
    ],
)
def test_proposal_build_missing_record_leaves_no_outputs(
    tmp_path: Path,
    capsys,
    customer_id: str,
    job_id: str,
    expected_error: str,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(
        _proposal_build_args(
            db_path,
            output_json,
            output_docx,
            customer_id=customer_id,
            job_id=job_id,
        )
    ) == 1
    assert expected_error in capsys.readouterr().err
    _assert_no_outputs(output_json, output_docx)


def test_proposal_build_customer_job_mismatch_uses_existing_composition(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    assert main(
        ["records", "import", "customer", str(CUSTOMER_EXAMPLE), "--db", str(db_path)]
    ) == 0
    assert main(
        ["records", "import", "job", str(SAMPLE_JOB_EXAMPLE), "--db", str(db_path)]
    ) == 0
    capsys.readouterr()
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(
        _proposal_build_args(
            db_path,
            output_json,
            output_docx,
            job_id="job-sample-north-prairie",
        )
    ) == 1
    assert "failed to compose proposal input" in capsys.readouterr().err
    _assert_no_outputs(output_json, output_docx)


@pytest.mark.parametrize("details_kind", ["missing", "directory", "malformed", "invalid"])
def test_proposal_build_invalid_details_leave_no_outputs(
    tmp_path: Path,
    capsys,
    details_kind: str,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    details_path = tmp_path / "details.json"
    if details_kind == "directory":
        details_path.mkdir()
    elif details_kind == "malformed":
        details_path.write_text("{", encoding="utf-8")
    elif details_kind == "invalid":
        details_path.write_text("{}", encoding="utf-8")
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(
        _proposal_build_args(
            db_path,
            output_json,
            output_docx,
            details_path=details_path,
        )
    ) == 1
    _assert_no_outputs(output_json, output_docx)


def test_proposal_build_unresolved_placeholders_leave_no_outputs(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    details_path = _write_details(
        tmp_path,
        lambda payload: payload.__setitem__(
            "item_description",
            "TODO: replace with explicit work",
        ),
    )
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(
        _proposal_build_args(
            db_path,
            output_json,
            output_docx,
            details_path=details_path,
        )
    ) == 1
    captured = capsys.readouterr()
    assert "unresolved placeholder text" in captured.err
    assert "item_description" in captured.err
    _assert_no_outputs(output_json, output_docx)


@pytest.mark.parametrize("template_kind", ["missing", "directory", "corrupt"])
def test_proposal_build_invalid_template_leaves_no_outputs(
    tmp_path: Path,
    capsys,
    template_kind: str,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    template_path = tmp_path / "template.docx"
    if template_kind == "directory":
        template_path.mkdir()
    elif template_kind == "corrupt":
        template_path.write_bytes(b"not a docx")
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(
        _proposal_build_args(
            db_path,
            output_json,
            output_docx,
            template_path=template_path,
        )
    ) == 1
    _assert_no_outputs(output_json, output_docx)


@pytest.mark.parametrize("blocked_output", ["json", "docx"])
def test_proposal_build_existing_output_blocks_both_artifacts_without_modification(
    tmp_path: Path,
    capsys,
    blocked_output: str,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"
    existing_path = output_json if blocked_output == "json" else output_docx
    existing_bytes = b"existing-output"
    existing_path.write_bytes(existing_bytes)

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    capsys.readouterr()

    assert existing_path.read_bytes() == existing_bytes
    other_path = output_docx if blocked_output == "json" else output_json
    assert not other_path.exists()


def test_proposal_build_rejects_identical_output_paths(tmp_path: Path, capsys) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_path = tmp_path / "proposal.output"

    assert main(_proposal_build_args(db_path, output_path, output_path)) == 1
    assert "output paths must be different" in capsys.readouterr().err
    assert not output_path.exists()


def test_proposal_build_json_staging_failure_cleans_outputs_and_temporaries(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_dir = tmp_path / "proposal"
    output_json = output_dir / "proposal.json"
    output_docx = output_dir / "proposal.docx"

    def fail_json_staging(*_args, **_kwargs):
        raise OSError("forced JSON staging failure")

    monkeypatch.setattr(
        proposal_build,
        "_write_proposal_input_json",
        fail_json_staging,
    )
    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    assert "failed to stage proposal input JSON" in capsys.readouterr().err
    _assert_no_outputs(output_json, output_docx)
    assert list(output_dir.iterdir()) == []


def test_proposal_build_renderer_failure_cleans_outputs_and_temporaries(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_dir = tmp_path / "proposal"
    output_json = output_dir / "proposal.json"
    output_docx = output_dir / "proposal.docx"

    def fail_render(*_args, **_kwargs):
        raise OSError("forced renderer failure")

    monkeypatch.setattr(proposal_build.DocxProposalRenderer, "render", fail_render)
    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    assert "failed to render proposal DOCX" in capsys.readouterr().err
    _assert_no_outputs(output_json, output_docx)
    assert list(output_dir.iterdir()) == []


def test_proposal_build_second_publication_failure_removes_first_output(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_dir = tmp_path / "proposal"
    output_json = output_dir / "proposal.json"
    output_docx = output_dir / "proposal.docx"
    original_publish = proposal_build._publish_staged_artifact
    call_count = 0

    def fail_second_publish(staged_path, final_path, created_final_paths):
        nonlocal call_count
        call_count += 1
        if call_count == 2:
            raise OSError("forced second publication failure")
        original_publish(staged_path, final_path, created_final_paths)

    monkeypatch.setattr(
        proposal_build,
        "_publish_staged_artifact",
        fail_second_publish,
    )
    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    assert "failed to publish final proposal outputs" in capsys.readouterr().err
    _assert_no_outputs(output_json, output_docx)
    assert list(output_dir.iterdir()) == []


def test_proposal_build_uses_noninitializing_read_only_repositories(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    def unexpected_initializer(*_args, **_kwargs):
        raise AssertionError("proposal-build must not initialize the records database")

    monkeypatch.setattr(cli, "create_sqlite_record_store", unexpected_initializer)
    monkeypatch.setattr(
        sqlite_records,
        "initialize_records_database",
        unexpected_initializer,
    )

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 0
    assert output_json.exists()
    assert output_docx.exists()


def test_existing_record_import_still_initializes_new_database(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"

    assert main(
        ["records", "import", "customer", str(CUSTOMER_EXAMPLE), "--db", str(db_path)]
    ) == 0
    capsys.readouterr()

    snapshot = _database_snapshot(db_path)
    sqlite_master = snapshot[2]
    table_names = {row[1] for row in sqlite_master if row[0] == "table"}
    assert {"customers", "jobs"}.issubset(table_names)


def test_proposal_build_leaves_zero_byte_database_unchanged(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    db_path.write_bytes(b"")
    before = _database_snapshot(db_path)
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    assert "failed to read records database" in capsys.readouterr().err

    assert _database_snapshot(db_path) == before
    assert before[1] == 0
    assert before[2] == ()
    _assert_no_outputs(output_json, output_docx)
    _assert_database_has_no_sidecars_or_outputs(tmp_path, db_path)


@pytest.mark.parametrize(
    ("missing_table", "preserved_table"),
    [("jobs", "customers"), ("customers", "jobs")],
)
def test_proposal_build_leaves_partial_database_unchanged(
    tmp_path: Path,
    capsys,
    missing_table: str,
    preserved_table: str,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    with sqlite3.connect(db_path) as connection:
        connection.execute(f"DROP TABLE {missing_table}")
        connection.commit()
    before = _database_snapshot(db_path)
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
    assert "failed to read records database" in capsys.readouterr().err

    after = _database_snapshot(db_path)
    assert after == before
    table_names = {row[1] for row in after[2] if row[0] == "table"}
    assert missing_table not in table_names
    assert preserved_table in table_names
    preserved_contents = dict(after[3])[preserved_table]
    assert len(preserved_contents) == 1
    _assert_no_outputs(output_json, output_docx)
    _assert_database_has_no_sidecars_or_outputs(tmp_path, db_path)


def test_read_only_repositories_read_records_and_sqlite_blocks_saves(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    before = _database_snapshot(db_path)
    customers = SQLiteCustomerRepository(db_path, initialize=False, read_only=True)
    jobs = SQLiteJobRepository(db_path, initialize=False, read_only=True)

    customer = customers.get_customer("customer-abby-hill")
    job = jobs.get_job("job-abby-hill")
    assert customer is not None
    assert job is not None
    assert customers.read_only is True
    assert jobs.read_only is True

    with pytest.raises(sqlite3.OperationalError, match="readonly|read-only"):
        customers.save_customer(customer)
    with pytest.raises(sqlite3.OperationalError, match="readonly|read-only"):
        jobs.save_job(job)

    assert _database_snapshot(db_path) == before
    _assert_database_has_no_sidecars_or_outputs(tmp_path, db_path)


@pytest.mark.parametrize(
    "repository_type",
    [SQLiteCustomerRepository, SQLiteJobRepository],
)
def test_read_only_repository_rejects_initialization(
    tmp_path: Path,
    repository_type,
) -> None:
    db_path = tmp_path / "records.sqlite"

    with pytest.raises(ValueError, match="read-only.*cannot initialize"):
        repository_type(db_path, read_only=True, initialize=True)

    assert not db_path.exists()


def test_immutable_read_only_repositories_do_not_create_wal_sidecars(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    _enable_persistent_wal_mode(db_path)
    before = _database_snapshot(db_path)

    customers = SQLiteCustomerRepository(db_path, initialize=False, read_only=True)
    jobs = SQLiteJobRepository(db_path, initialize=False, read_only=True)
    customer = customers.get_customer("customer-abby-hill")
    job = jobs.get_job("job-abby-hill")

    assert customer is not None
    assert customer.display_name == "Abby Hill"
    assert job is not None
    assert job.customer_id == customer.customer_id
    assert _database_snapshot(db_path) == before
    _assert_no_sqlite_sidecars(db_path)
    _assert_database_has_no_sidecars_or_outputs(tmp_path, db_path)


def test_proposal_build_leaves_closed_wal_database_and_sidecars_unchanged(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    _enable_persistent_wal_mode(db_path)
    before = _database_snapshot(db_path)
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"

    assert main(_proposal_build_args(db_path, output_json, output_docx)) == 0
    captured = capsys.readouterr()

    proposal = ProposalInput.model_validate_json(output_json.read_text(encoding="utf-8"))
    assert proposal.customer_name == "Abby Hill"
    assert proposal.street_address == "123 Main St."
    assert proposal.pricing.amount == Decimal("3000.00")
    assert proposal.pricing.is_starting_at is True
    assert proposal.notes == []
    assert output_docx.stat().st_size > 0
    docx_text = _docx_text(output_docx)
    assert "Abby Hill" in docx_text
    assert "123 Main St." in docx_text
    assert "Starting at $3,000.00" in docx_text
    assert f"ProposalInput JSON: {output_json}" in captured.out
    assert f"Proposal DOCX: {output_docx}" in captured.out

    assert _database_snapshot(db_path) == before
    _assert_no_sqlite_sidecars(db_path)
    assert set(tmp_path.iterdir()) == {db_path, output_json, output_docx}


def test_immutable_reads_reject_existing_wal_state_without_mutation(
    tmp_path: Path,
    capsys,
) -> None:
    db_path = tmp_path / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    writer = sqlite3.connect(db_path)
    try:
        result = writer.execute("PRAGMA journal_mode=WAL").fetchone()
        assert result is not None
        assert result[0].lower() == "wal"
        writer.execute("PRAGMA wal_autocheckpoint=0")
        writer.execute(
            "UPDATE customers SET display_name = ? WHERE customer_id = ?",
            ("Committed only in WAL", "customer-abby-hill"),
        )
        writer.commit()

        before_database = db_path.read_bytes()
        before_sidecars = _sqlite_sidecar_snapshot(db_path)
        assert before_sidecars["-wal"] is not None
        assert before_sidecars["-shm"] is not None
        assert before_sidecars["-journal"] is None

        customers = SQLiteCustomerRepository(db_path, initialize=False, read_only=True)
        jobs = SQLiteJobRepository(db_path, initialize=False, read_only=True)
        with pytest.raises(sqlite3.OperationalError, match="without sidecar files"):
            customers.get_customer("customer-abby-hill")
        with pytest.raises(sqlite3.OperationalError, match="without sidecar files"):
            jobs.get_job("job-abby-hill")

        output_json = tmp_path / "proposal.json"
        output_docx = tmp_path / "proposal.docx"
        assert main(_proposal_build_args(db_path, output_json, output_docx)) == 1
        assert "failed to read records database" in capsys.readouterr().err
        _assert_no_outputs(output_json, output_docx)

        assert db_path.read_bytes() == before_database
        assert _sqlite_sidecar_snapshot(db_path) == before_sidecars
    finally:
        writer.close()


def test_immutable_reads_check_sidecars_beside_resolved_database(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    database_dir = tmp_path / "database"
    database_dir.mkdir()
    db_path = database_dir / "records.sqlite"
    _import_abby_records(db_path)
    capsys.readouterr()
    link_dir = tmp_path / "link"
    link_dir.mkdir()
    linked_db_path = link_dir / "records.sqlite"
    simulated_symlink = False
    try:
        linked_db_path.symlink_to(db_path)
    except OSError:
        simulated_symlink = True
        linked_db_path.write_bytes(b"simulated-symlink")
        original_resolve = Path.resolve

        def resolve_simulated_symlink(path: Path, *args, **kwargs) -> Path:
            if path == linked_db_path:
                return db_path
            return original_resolve(path, *args, **kwargs)

        monkeypatch.setattr(Path, "resolve", resolve_simulated_symlink)

    writer = sqlite3.connect(db_path)
    try:
        result = writer.execute("PRAGMA journal_mode=WAL").fetchone()
        assert result is not None
        assert result[0].lower() == "wal"
        writer.execute("PRAGMA wal_autocheckpoint=0")
        writer.execute(
            "UPDATE customers SET display_name = ? WHERE customer_id = ?",
            ("Committed only in target WAL", "customer-abby-hill"),
        )
        writer.commit()
        before_database = db_path.read_bytes()
        before_sidecars = _sqlite_sidecar_snapshot(db_path)
        before_link = linked_db_path.read_bytes() if simulated_symlink else None
        assert before_sidecars["-wal"] is not None
        assert before_sidecars["-shm"] is not None

        customers = SQLiteCustomerRepository(
            linked_db_path,
            initialize=False,
            read_only=True,
        )
        jobs = SQLiteJobRepository(linked_db_path, initialize=False, read_only=True)
        with pytest.raises(sqlite3.OperationalError, match="without sidecar files"):
            customers.get_customer("customer-abby-hill")
        with pytest.raises(sqlite3.OperationalError, match="without sidecar files"):
            jobs.get_job("job-abby-hill")

        output_json = tmp_path / "proposal.json"
        output_docx = tmp_path / "proposal.docx"
        assert main(_proposal_build_args(linked_db_path, output_json, output_docx)) == 1
        assert "failed to read records database" in capsys.readouterr().err
        _assert_no_outputs(output_json, output_docx)

        assert db_path.read_bytes() == before_database
        assert _sqlite_sidecar_snapshot(db_path) == before_sidecars
        if simulated_symlink:
            assert linked_db_path.read_bytes() == before_link
        _assert_no_sqlite_sidecars(linked_db_path)
    finally:
        writer.close()
