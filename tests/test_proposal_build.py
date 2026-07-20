"""Focused tests for the read-only in-process proposal build service."""

from __future__ import annotations

import hashlib
import sqlite3
import subprocess
from pathlib import Path

import pytest
from docx import Document

import phoenix_office.proposal_build as proposal_build
import phoenix_office.records.sqlite as sqlite_records
from phoenix_office.models.proposal import ProposalInput
from phoenix_office.records import (
    create_sqlite_record_store,
    customer_record_from_json_file,
    job_record_from_json,
    record_proposal_details_from_file,
)

ROOT = Path(__file__).parents[1]
CUSTOMER_EXAMPLE = ROOT / "examples" / "records" / "customer_abby_hill.json"
JOB_EXAMPLE = ROOT / "examples" / "records" / "job_abby_hill.json"
DETAILS_EXAMPLE = ROOT / "examples" / "records" / "proposal_details_abby_hill.json"
TEMPLATE = ROOT / "tests" / "fixtures" / "templates" / "a1_proposal_template.docx"
SAMPLE_JOB_EXAMPLE = ROOT / "examples" / "records" / "job_sample_north_prairie.json"


def _save_records(
    database_path: Path,
    *,
    save_customer: bool = True,
    save_job: bool = True,
    job_path: Path = JOB_EXAMPLE,
) -> None:
    store = create_sqlite_record_store(database_path)
    if save_customer:
        store.customers.save_customer(customer_record_from_json_file(CUSTOMER_EXAMPLE))
    if save_job:
        store.jobs.save_job(job_record_from_json(job_path.read_text(encoding="utf-8")))


def _request(
    tmp_path: Path,
    database_path: Path,
    *,
    customer_id: str = "customer-abby-hill",
    job_id: str = "job-abby-hill",
    details=None,
    template_path: Path = TEMPLATE,
    output_json: Path | None = None,
    output_docx: Path | None = None,
) -> proposal_build.ProposalDraftBuildRequest:
    return proposal_build.ProposalDraftBuildRequest(
        customer_id=customer_id,
        job_id=job_id,
        details=details or record_proposal_details_from_file(DETAILS_EXAMPLE),
        database_path=database_path,
        template_path=template_path,
        proposal_input_json_output_path=output_json or tmp_path / "output" / "proposal.json",
        proposal_docx_output_path=output_docx or tmp_path / "output" / "proposal.docx",
    )


def _database_snapshot(database_path: Path) -> tuple[object, ...]:
    database_uri = f"{database_path.resolve().as_uri()}?mode=ro&immutable=1"
    with sqlite3.connect(database_uri, uri=True) as connection:
        sqlite_master = tuple(
            connection.execute(
                """
                SELECT type, name, tbl_name, sql
                FROM sqlite_master
                ORDER BY type, name, tbl_name
                """
            ).fetchall()
        )
        table_names = {
            row[1]
            for row in sqlite_master
            if row[0] == "table" and row[1] != "sqlite_sequence"
        }
        table_contents = tuple(
            (table_name, tuple(connection.execute(f"SELECT * FROM {table_name} ORDER BY rowid")))
            for table_name in sorted(table_names)
        )
    database_bytes = database_path.read_bytes()
    return (
        hashlib.sha256(database_bytes).hexdigest(),
        len(database_bytes),
        sqlite_master,
        table_contents,
    )


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    return " ".join(" ".join(values).split())


def _assert_no_outputs_or_temporaries(
    tmp_path: Path,
    output_json: Path,
    output_docx: Path,
) -> None:
    assert not output_json.exists()
    assert not output_docx.exists()
    assert list(tmp_path.rglob("*.tmp")) == []


def _failure_lines(
    request: proposal_build.ProposalDraftBuildRequest,
) -> tuple[str, ...]:
    with pytest.raises(proposal_build._ProposalDraftBuildFailure) as captured:
        proposal_build.build_proposal_draft(request)
    return captured.value.stderr_lines


def test_build_proposal_draft_returns_published_reviewable_artifacts_without_mutation(
    tmp_path: Path,
    capsys,
    monkeypatch,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    before = _database_snapshot(database_path)
    request = _request(tmp_path, database_path)

    def unexpected_call(*_args, **_kwargs):
        raise AssertionError("out-of-scope operation invoked")

    monkeypatch.setattr(
        sqlite_records,
        "initialize_records_database",
        unexpected_call,
    )
    monkeypatch.setattr(
        proposal_build.SQLiteCustomerRepository,
        "save_customer",
        unexpected_call,
    )
    monkeypatch.setattr(
        proposal_build.SQLiteJobRepository,
        "save_job",
        unexpected_call,
    )
    for name in (
        "delete_customer",
        "delete_job",
        "import_customer_record_file",
        "import_job_record_file",
        "execute_core_command",
        "execute_sdk_command",
        "send_email",
        "upload_file",
        "file_proposal",
    ):
        monkeypatch.setattr(proposal_build, name, unexpected_call, raising=False)
    monkeypatch.setattr(subprocess, "run", unexpected_call)
    monkeypatch.setattr(subprocess, "Popen", unexpected_call)
    original_select = proposal_build._select_existing_proposal_records
    record_preflights: list[tuple[Path, str, str]] = []

    def select_records(
        *,
        database_path: Path,
        customer_id: str,
        job_id: str,
    ):
        record_preflights.append((database_path, customer_id, job_id))
        return original_select(
            database_path=database_path,
            customer_id=customer_id,
            job_id=job_id,
        )

    monkeypatch.setattr(
        proposal_build,
        "_select_existing_proposal_records",
        select_records,
    )

    result = proposal_build.build_proposal_draft(request)
    captured = capsys.readouterr()

    assert isinstance(result, proposal_build.ProposalDraftBuildResult)
    assert result.proposal_input_json_path == request.proposal_input_json_output_path
    assert result.proposal_docx_path == request.proposal_docx_output_path
    assert result.proposal_input_json_path.exists()
    assert result.proposal_docx_path.exists()
    assert result.proposal_docx_path.stat().st_size > 0
    assert ProposalInput.model_validate_json(
        result.proposal_input_json_path.read_text(encoding="utf-8")
    ) == result.proposal_input
    assert "Abby Hill" in _docx_text(result.proposal_docx_path)
    assert "Starting at $3,000.00" in _docx_text(result.proposal_docx_path)
    assert result.summary_lines == (
        "Customer: Abby Hill",
        "Site Address: 123 Main St., Menomonee Falls, WI 53051",
        "Item Description: Removal of 1,000 Gallon Aboveground Storage Tank",
        "Scope Items: 4",
        "Pricing Lines: 1",
        "Total: Starting at $3,000.00",
        "Notes: none",
        "Company: A-1 Tank Removal LLC",
    )
    assert _database_snapshot(database_path) == before
    assert record_preflights == [
        (database_path, "customer-abby-hill", "job-abby-hill")
    ]
    for suffix in ("-wal", "-shm", "-journal"):
        assert not Path(f"{database_path}{suffix}").exists()
    assert captured.out == ""
    assert captured.err == ""


def test_build_proposal_draft_rejects_identical_output_paths(tmp_path: Path) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    output_path = tmp_path / "proposal.output"

    assert _failure_lines(
        _request(
            tmp_path,
            database_path,
            output_json=output_path,
            output_docx=output_path,
        )
    ) == ("Error: proposal output paths must be different",)
    _assert_no_outputs_or_temporaries(tmp_path, output_path, output_path)


@pytest.mark.parametrize("existing_output", ["json", "docx", "both"])
def test_build_proposal_draft_refuses_existing_outputs_without_modification(
    tmp_path: Path,
    existing_output: str,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    output_json = tmp_path / "proposal.json"
    output_docx = tmp_path / "proposal.docx"
    existing_paths = []
    if existing_output in {"json", "both"}:
        existing_paths.append(output_json)
    if existing_output in {"docx", "both"}:
        existing_paths.append(output_docx)
    for path in existing_paths:
        path.write_bytes(b"existing-output")

    lines = _failure_lines(
        _request(
            tmp_path,
            database_path,
            output_json=output_json,
            output_docx=output_docx,
        )
    )

    assert len(lines) == len(existing_paths)
    assert all("output already exists" in line for line in lines)
    for path in existing_paths:
        assert path.read_bytes() == b"existing-output"
    for path in {output_json, output_docx} - set(existing_paths):
        assert not path.exists()
    assert list(tmp_path.rglob("*.tmp")) == []


@pytest.mark.parametrize("database_kind", ["missing", "directory"])
def test_build_proposal_draft_rejects_missing_or_non_file_database(
    tmp_path: Path,
    database_kind: str,
) -> None:
    database_path = tmp_path / "records.sqlite"
    if database_kind == "directory":
        database_path.mkdir()
    request = _request(tmp_path, database_path)

    lines = _failure_lines(request)

    expected = "does not exist" if database_kind == "missing" else "path is not a file"
    assert expected in lines[0]
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


@pytest.mark.parametrize("database_kind", ["zero-byte", "invalid", "partial"])
def test_build_proposal_draft_rejects_unusable_database_without_mutation(
    tmp_path: Path,
    database_kind: str,
) -> None:
    database_path = tmp_path / "records.sqlite"
    if database_kind == "zero-byte":
        database_path.write_bytes(b"")
    elif database_kind == "invalid":
        database_path.write_bytes(b"not a sqlite database")
    else:
        _save_records(database_path)
        with sqlite3.connect(database_path) as connection:
            connection.execute("DROP TABLE jobs")
            connection.commit()
    before = database_path.read_bytes()
    request = _request(tmp_path, database_path)

    assert "failed to read records database" in _failure_lines(request)[0]
    assert database_path.read_bytes() == before
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_rejects_existing_sqlite_sidecar_without_mutation(
    tmp_path: Path,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    database_before = database_path.read_bytes()
    sidecar_path = Path(f"{database_path}-wal")
    sidecar_path.write_bytes(b"existing-sidecar")
    request = _request(tmp_path, database_path)

    assert "failed to read records database" in _failure_lines(request)[0]
    assert database_path.read_bytes() == database_before
    assert sidecar_path.read_bytes() == b"existing-sidecar"
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


@pytest.mark.parametrize(
    ("save_customer", "save_job", "customer_id", "job_id", "expected"),
    [
        (False, True, "missing-customer", "job-abby-hill", "Customer not found"),
        (True, False, "customer-abby-hill", "missing-job", "Job not found"),
    ],
)
def test_build_proposal_draft_rejects_missing_saved_record(
    tmp_path: Path,
    save_customer: bool,
    save_job: bool,
    customer_id: str,
    job_id: str,
    expected: str,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(
        database_path,
        save_customer=save_customer,
        save_job=save_job,
    )
    request = _request(
        tmp_path,
        database_path,
        customer_id=customer_id,
        job_id=job_id,
    )

    assert expected in _failure_lines(request)[0]
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_rejects_customer_job_mismatch(tmp_path: Path) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path, job_path=SAMPLE_JOB_EXAMPLE)
    request = _request(
        tmp_path,
        database_path,
        job_id="job-sample-north-prairie",
    )

    assert "failed to compose proposal input" in _failure_lines(request)[0]
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


@pytest.mark.parametrize("template_kind", ["missing", "directory", "corrupt"])
def test_build_proposal_draft_rejects_unusable_template(
    tmp_path: Path,
    template_kind: str,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    template_path = tmp_path / "template.docx"
    if template_kind == "directory":
        template_path.mkdir()
    elif template_kind == "corrupt":
        template_path.write_bytes(b"not a docx")
    request = _request(tmp_path, database_path, template_path=template_path)

    lines = _failure_lines(request)

    assert "DOCX template" in lines[0]
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_rejects_unresolved_placeholders(tmp_path: Path) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    details = record_proposal_details_from_file(DETAILS_EXAMPLE).model_copy(
        update={"item_description": "TODO: replace with explicit work"}
    )
    request = _request(tmp_path, database_path, details=details)

    lines = _failure_lines(request)

    assert lines[0] == (
        "Error: unresolved placeholder text in composed proposal; "
        "refusing proposal build."
    )
    assert lines[1] == "Placeholder fields: item_description"
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_output_directory_failure_leaves_no_artifacts(
    tmp_path: Path,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    blocked_parent = tmp_path / "blocked"
    blocked_parent.write_bytes(b"not a directory")
    output_json = blocked_parent / "proposal.json"
    output_docx = blocked_parent / "proposal.docx"
    request = _request(
        tmp_path,
        database_path,
        output_json=output_json,
        output_docx=output_docx,
    )

    assert _failure_lines(request) == (
        "Error: failed to prepare proposal output directories",
    )
    _assert_no_outputs_or_temporaries(tmp_path, output_json, output_docx)


def test_build_proposal_draft_json_staging_failure_cleans_temporary_sibling(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    request = _request(tmp_path, database_path)

    def fail_json_staging(*_args, **_kwargs):
        raise OSError("forced JSON staging failure")

    monkeypatch.setattr(
        proposal_build,
        "_write_proposal_input_json",
        fail_json_staging,
    )

    assert "failed to stage proposal input JSON" in _failure_lines(request)[0]
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


@pytest.mark.parametrize("render_failure", ["exception", "empty"])
def test_build_proposal_draft_render_failure_cleans_staged_artifacts(
    tmp_path: Path,
    monkeypatch,
    render_failure: str,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    request = _request(tmp_path, database_path)

    def fail_or_leave_empty(*_args, **_kwargs):
        if render_failure == "exception":
            raise OSError("forced renderer failure")
        return request.proposal_docx_output_path

    monkeypatch.setattr(
        proposal_build.DocxProposalRenderer,
        "render",
        fail_or_leave_empty,
    )

    assert "failed to render proposal DOCX" in _failure_lines(request)[0]
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_missing_staged_artifact_cleans_remaining_sibling(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    request = _request(tmp_path, database_path)
    original_write = proposal_build._write_proposal_input_json

    def remove_staged_json(proposal, output_path):
        original_write(proposal, output_path)
        output_path.unlink()

    monkeypatch.setattr(
        proposal_build,
        "_write_proposal_input_json",
        remove_staged_json,
    )

    assert _failure_lines(request) == (
        "Error: proposal output staging did not produce both artifacts",
    )
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_partial_publication_failure_rolls_back_first_output(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    request = _request(tmp_path, database_path)
    original_publish = proposal_build._publish_staged_artifact
    call_count = 0

    def fail_second_publish(staged_path, final_path, created_final_paths):
        nonlocal call_count
        call_count += 1
        if call_count == 2:
            raise OSError("forced second publication failure")
        original_publish(staged_path, final_path, created_final_paths)

    monkeypatch.setattr(
        proposal_build,
        "_publish_staged_artifact",
        fail_second_publish,
    )

    assert _failure_lines(request) == (
        "Error: failed to publish final proposal outputs",
    )
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )


def test_build_proposal_draft_incomplete_final_output_verification_rolls_back(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database_path = tmp_path / "records.sqlite"
    _save_records(database_path)
    request = _request(tmp_path, database_path)
    original_publish = proposal_build._publish_staged_artifact

    def publish_empty_docx(staged_path, final_path, created_final_paths):
        if final_path == request.proposal_docx_output_path:
            with final_path.open("xb"):
                created_final_paths.append(final_path)
            return
        original_publish(staged_path, final_path, created_final_paths)

    monkeypatch.setattr(
        proposal_build,
        "_publish_staged_artifact",
        publish_empty_docx,
    )

    assert _failure_lines(request) == (
        "Error: final proposal output verification failed",
    )
    _assert_no_outputs_or_temporaries(
        tmp_path,
        request.proposal_input_json_output_path,
        request.proposal_docx_output_path,
    )
